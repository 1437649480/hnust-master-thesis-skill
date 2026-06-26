# -*- coding: utf-8 -*-
"""
湖南科技大学硕士学位论文生成脚本
根据《湖南科技大学研究生学位论文撰写规范》生成符合格式要求的示例docx文件
用法: python generate_thesis.py [输出路径，默认: 硕士论文格式示例.docx]
"""

import sys, os
from docx import Document
from docx.shared import Pt, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── 常量 ───
SPEC = {
    'page': {
        'width': Cm(21), 'height': Cm(29.7),  # A4
        'top': Cm(3), 'bottom': Cm(3),
        'left': Cm(2.5), 'right': Cm(2.5),
        'header': Cm(2), 'footer': Cm(2),
    },
}


# ─── 工具函数 ───

def set_page(section):
    """设置页面大小和边距"""
    section.page_width = SPEC['page']['width']
    section.page_height = SPEC['page']['height']
    section.top_margin = SPEC['page']['top']
    section.bottom_margin = SPEC['page']['bottom']
    section.left_margin = SPEC['page']['left']
    section.right_margin = SPEC['page']['right']
    section.header_distance = SPEC['page']['header']
    section.footer_distance = SPEC['page']['footer']


def set_run_font(run, font_name, size, bold=False, east_asian=None):
    """设置run字体，同时设置中西文字体"""
    run.font.name = font_name
    run.font.size = size
    run.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), east_asian or font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)


def add_paragraph(doc, text, font_name='宋体', size=Pt(12), bold=False, align=None,
                  space_before=0, space_after=0, line_spacing=1.25, first_line_indent=None,
                  font_name_east=None):
    """添加格式化段落"""
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    # 处理段前段后间距（使用EMU单位）
    # 1pt = 12700 EMU
    if isinstance(space_before, (int, float)):
        pf.space_before = int(space_before * 12700)
    elif hasattr(space_before, '__int__'):
        pf.space_before = int(space_before)
    if isinstance(space_after, (int, float)):
        pf.space_after = int(space_after * 12700)
    elif hasattr(space_after, '__int__'):
        pf.space_after = int(space_after)
    pf.line_spacing = line_spacing
    if first_line_indent:
        pf.first_line_indent = first_line_indent
    run = p.add_run(text)
    set_run_font(run, font_name, size, bold, font_name_east)
    return p


def add_title_h1(doc, text):
    """章标题：小2号宋体加粗(18pt)，上下各空一行，单倍行距"""
    return add_paragraph(doc, text, '宋体', Pt(18), True,
                         WD_ALIGN_PARAGRAPH.CENTER, 12, 12, 1.0)


def add_title_h2(doc, text):
    """节标题：小3号宋体加粗(15pt)，上下各空0.5行，单倍行距"""
    return add_paragraph(doc, text, '宋体', Pt(15), True,
                         WD_ALIGN_PARAGRAPH.LEFT, 6, 6, 1.0)


def add_title_h3(doc, text):
    """二级标题：4号宋体加粗(14pt)，上下各空0.5行，单倍行距"""
    return add_paragraph(doc, text, '宋体', Pt(14), True,
                         WD_ALIGN_PARAGRAPH.LEFT, 6, 6, 1.0)


def add_body(doc, text):
    """正文段落：小4号宋体(12pt)，1.25倍行距，首行缩进2字符"""
    return add_paragraph(doc, text, '宋体', Pt(12), False,
                         WD_ALIGN_PARAGRAPH.LEFT, 0, 0, 1.25, Pt(24))


def add_ref(doc, text):
    """参考文献条目：小5号宋体(9pt)"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.25
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.name = 'Times New Roman'
    run.bold = False
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), '宋体')
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    return p


def add_page_break(doc):
    """添加分页符"""
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._element.append(br)


def setup_header(section, text='湖南科技大学硕士学位论文'):
    """设置页眉：五号宋体居中"""
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.clear()
    run = p.add_run(text)
    set_run_font(run, '宋体', Pt(10.5), False)


def setup_footer(section):
    """设置页脚：-X- 格式，5号Times New Roman居中"""
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.clear()

    r1 = p.add_run('-')
    set_run_font(r1, 'Times New Roman', Pt(10.5))

    r2 = p.add_run()
    set_run_font(r2, 'Times New Roman', Pt(10.5))
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    r2._element.append(fld_begin)
    instr = OxmlElement('w:instrText')
    instr.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    instr.text = ' PAGE '
    r2._element.append(instr)
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    r2._element.append(fld_end)

    r3 = p.add_run('-')
    set_run_font(r3, 'Times New Roman', Pt(10.5))


def add_bilingual_caption(doc, cn_text, en_text):
    """添加中英文双语图表标题"""
    add_paragraph(doc, cn_text, '宋体', Pt(12), True, WD_ALIGN_PARAGRAPH.CENTER, 6, 0)
    add_paragraph(doc, en_text, 'Times New Roman', Pt(12), True, WD_ALIGN_PARAGRAPH.CENTER, 0, 6)


def add_figure_placeholder(doc, cn_label, en_label):
    """添加图表占位"""
    add_bilingual_caption(doc, cn_label, en_label)
    add_paragraph(doc, '[此处插入图片]', '宋体', Pt(10.5), False, WD_ALIGN_PARAGRAPH.CENTER, 0, 0)


# ─── 主生成函数 ───

def generate(filepath):
    """生成完整硕士论文示例文档"""
    doc = Document()

    # ── 页面设置 ──
    section = doc.sections[0]
    set_page(section)
    setup_header(section)
    setup_footer(section)

    # ══════════════════════════════════════════════
    # 一、扉页/题名页
    # ══════════════════════════════════════════════
    add_paragraph(doc, '', '宋体', Pt(12), False, WD_ALIGN_PARAGRAPH.CENTER, 0, 0, 1.0)
    add_paragraph(doc, '湖 南 科 技 大 学', '宋体', Pt(26), True, WD_ALIGN_PARAGRAPH.CENTER, 0, 6)
    add_paragraph(doc, '硕 士 学 位 论 文', '宋体', Pt(26), True, WD_ALIGN_PARAGRAPH.CENTER, 0, 24)

    add_paragraph(doc, '', '宋体', Pt(12), False, WD_ALIGN_PARAGRAPH.CENTER, 0, 0, 1.0)

    cover_data = [
        ('题　目', ''),
        ('学　科　专　业', ''),
        ('研　究　方　向', ''),
        ('研　究　生', ''),
        ('指导教师', ''),
        ('论文编号', ''),
    ]
    table = doc.add_table(rows=len(cover_data), cols=2)
    for i, (label, _) in enumerate(cover_data):
        row = table.rows[i]
        cell0 = row.cells[0]
        cell0.text = ''
        p0 = cell0.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run0 = p0.add_run(label)
        set_run_font(run0, '宋体', Pt(16), True)
        cell1 = row.cells[1]
        cell1.text = ''

    add_paragraph(doc, '', '宋体', Pt(12), False, WD_ALIGN_PARAGRAPH.CENTER, 12, 0, 1.0)
    add_paragraph(doc, '二〇二六年六月', '宋体', Pt(16), True, WD_ALIGN_PARAGRAPH.CENTER, 6, 0)

    add_page_break(doc)

    # ══════════════════════════════════════════════
    # 二、中文摘要
    # ══════════════════════════════════════════════
    add_paragraph(doc, '摘  要', '宋体', Pt(18), True, WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, '', '宋体', Pt(12), False, WD_ALIGN_PARAGRAPH.CENTER, 0, 0, 1.0)

    add_body(doc, '随着经济全球化的深入推进，外汇储备在国家宏观经济调控中的作用日益突出。我国自改革开放以来，外汇储备规模经历了从无到有、从小到大的发展过程，目前已成为全球最大的外汇储备持有国。充足的外汇储备有助于维护国际收支平衡、稳定汇率、抵御外部金融冲击，但过高的储备规模也带来了管理成本上升、货币供应被动扩张等问题。')
    add_body(doc, '本文首先梳理了我国外汇储备的发展历程，将其划分为规模较小、快速增长、稳步增长和大幅增长四个阶段。其次，深入分析了推动外汇储备增长的内外部因素，包括贸易顺差扩大、外商直接投资增加、国际资本流入以及汇率制度变革等。再次，运用协整分析和误差修正模型，探讨了外汇储备与主要宏观经济变量之间的长期均衡关系和短期波动特征。最后，在借鉴日本、新加坡、挪威等国成熟管理经验的基础上，提出了优化我国外汇储备管理的政策建议。')
    add_body(doc, '研究结果表明：我国外汇储备的增长与贸易顺差、外商直接投资存在显著的正向协整关系；短期内，人民币汇率预期和国际利差是影响储备变动的重要因素。本文建议从完善多元化投资机制、增强汇率弹性、推进人民币国际化等方面入手，提高外汇储备的使用效率和风险管理水平。')

    add_paragraph(doc, '', '宋体', Pt(12), False, WD_ALIGN_PARAGRAPH.CENTER, 0, 0, 1.0)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r1 = p.add_run('关键词：')
    set_run_font(r1, '黑体', Pt(12), True)
    r2 = p.add_run('外汇储备；宏观经济；协整分析；国际比较；管理策略')
    set_run_font(r2, '宋体', Pt(12))

    add_page_break(doc)

    # ══════════════════════════════════════════════
    # 三、英文摘要
    # ══════════════════════════════════════════════
    add_paragraph(doc, 'ABSTRACT', 'Times New Roman', Pt(18), True, WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, '', 'Times New Roman', Pt(12), False, WD_ALIGN_PARAGRAPH.CENTER, 0, 0, 1.0)

    add_paragraph(doc,
        'With the deepening of economic globalization, foreign exchange reserves have become increasingly '
        'prominent in national macroeconomic regulation and control. Since the reform and opening-up, '
        'China\'s foreign exchange reserves have experienced a development process from scratch to a '
        'large scale, and China has now become the world\'s largest holder of foreign exchange reserves. '
        'Sufficient foreign exchange reserves help maintain the balance of international payments, '
        'stabilize exchange rates, and resist external financial shocks. However, excessive reserves '
        'also bring about rising management costs and passive expansion of money supply.',
        'Times New Roman', Pt(12), False, WD_ALIGN_PARAGRAPH.LEFT, 0, 0, 1.25, Pt(24))

    add_paragraph(doc,
        'This paper first reviews the development history of China\'s foreign exchange reserves and '
        'divides it into four stages: small scale, rapid growth, steady growth, and substantial growth. '
        'Secondly, it analyzes the internal and external factors driving the growth of foreign exchange '
        'reserves, including the expansion of trade surplus, increase in foreign direct investment, '
        'inflow of international capital, and reforms in the exchange rate system. Thirdly, by employing '
        'cointegration analysis and error correction models, it explores the long-term equilibrium '
        'relationship and short-term fluctuation characteristics between foreign exchange reserves and '
        'major macroeconomic variables. Finally, based on the mature management experience of countries '
        'such as Japan, Singapore, and Norway, policy recommendations for optimizing China\'s foreign '
        'exchange reserve management are proposed.',
        'Times New Roman', Pt(12), False, WD_ALIGN_PARAGRAPH.LEFT, 0, 0, 1.25, Pt(24))

    add_paragraph(doc,
        'The research results show that the growth of China\'s foreign exchange reserves has a '
        'significant positive cointegration relationship with trade surplus and foreign direct '
        'investment. In the short term, RMB exchange rate expectations and international interest '
        'rate differentials are important factors affecting reserve changes. This paper suggests '
        'improving the efficiency of foreign exchange reserve use and risk management through '
        'diversifying investment mechanisms, enhancing exchange rate flexibility, and promoting '
        'RMB internationalization.',
        'Times New Roman', Pt(12), False, WD_ALIGN_PARAGRAPH.LEFT, 0, 0, 1.25, Pt(24))

    add_paragraph(doc, '', 'Times New Roman', Pt(12), False, WD_ALIGN_PARAGRAPH.CENTER, 0, 0, 1.0)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r1 = p.add_run('Keywords: ')
    set_run_font(r1, 'Times New Roman', Pt(12), True)
    r2 = p.add_run('foreign exchange reserves; macroeconomics; cointegration analysis; '
                   'international comparison; management strategy')
    set_run_font(r2, 'Times New Roman', Pt(12))

    add_page_break(doc)

    # ══════════════════════════════════════════════
    # 四、目录
    # ══════════════════════════════════════════════
    add_paragraph(doc, '目  录', '黑体', Pt(16), False, WD_ALIGN_PARAGRAPH.CENTER, 0, 12)

    toc_items = [
        ('摘  要', 'I', 0),
        ('ABSTRACT', 'II', 0),
        ('第一章 引言', '1', 0),
        ('1.1 研究背景', '1', 1),
        ('1.2 研究目的与意义', '2', 1),
        ('1.3 国内外研究现状', '3', 1),
        ('1.4 研究内容与方法', '5', 1),
        ('第二章 我国外汇储备的演变历程', '7', 0),
        ('2.1 规模较小阶段（1978—1993年）', '7', 1),
        ('2.2 快速增长阶段（1994—1997年）', '8', 1),
        ('2.3 稳步增长阶段（1998—2000年）', '9', 1),
        ('2.4 大幅增长阶段（2001年至今）', '10', 1),
        ('第三章 外汇储备影响因素实证分析', '13', 0),
        ('3.1 变量选取与数据说明', '13', 1),
        ('3.2 平稳性检验', '14', 1),
        ('3.3 协整检验与误差修正模型', '15', 1),
        ('3.4 实证结果分析', '17', 1),
        ('第四章 国际比较与经验借鉴', '19', 0),
        ('4.1 日本外汇储备管理', '19', 1),
        ('4.2 新加坡外汇储备管理', '20', 1),
        ('4.3 挪威外汇储备管理', '21', 1),
        ('4.4 经验总结与启示', '22', 1),
        ('第五章 结论与展望', '24', 0),
        ('5.1 主要结论', '24', 1),
        ('5.2 政策建议', '25', 1),
        ('5.3 研究不足与展望', '26', 1),
        ('参考文献', '27', 0),
        ('致  谢', '29', 0),
        ('附录A 攻读学位期间发表的论文', '30', 0),
    ]

    for item, page, level in toc_items:
        is_chapter = level == 0
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.tab_stops.add_tab_stop(Cm(14.5), alignment=WD_ALIGN_PARAGRAPH.RIGHT, leader=1)
        if level == 1:
            p.paragraph_format.left_indent = Cm(0.74)

        r1 = p.add_run(item)
        set_run_font(r1, '宋体', Pt(14) if is_chapter else Pt(12), is_chapter)
        r2 = p.add_run('\t')
        r3 = p.add_run(page)
        set_run_font(r3, 'Times New Roman', Pt(12))

    add_page_break(doc)

    # ══════════════════════════════════════════════
    # 五、正文：第一章 引言
    # ══════════════════════════════════════════════
    add_title_h1(doc, '第一章 引言')

    add_title_h2(doc, '1.1 研究背景')
    add_body(doc, '外汇储备是指由一国货币当局持有的、可自由兑换的外国货币资产，是国际储备的重要组成部分。它在维护国际收支平衡、稳定汇率、增强国际信用以及应对突发性金融危机等方面发挥着不可替代的作用。随着经济全球化进程的不断加快，外汇储备的规模和管理效率已成为衡量一个国家经济实力和金融安全的重要标志。')
    add_body(doc, '改革开放以来，我国经济持续高速增长，对外贸易规模不断扩大，外汇储备随之大幅增加。1978年，我国外汇储备仅有1.67亿美元；到1996年，首次突破1000亿美元；2006年，我国超越日本成为全球最大的外汇储备持有国；截至2024年末，我国外汇储备余额达到3.2万亿美元，约占全球外汇储备总额的四分之一。如此庞大的外汇储备规模，在为我国经济安全提供坚实保障的同时，也带来了储备管理效率、货币政策独立性以及资产保值增值等方面的挑战。')

    add_title_h2(doc, '1.2 研究目的与意义')
    add_body(doc, '本文旨在系统研究我国外汇储备的增长机理、影响因素及管理策略，为优化外汇储备管理提供理论支撑和政策参考。具体研究目的包括：第一，梳理我国外汇储备的发展演变历程，揭示其阶段性特征和内在规律；第二，运用计量经济学方法，定量分析影响外汇储备变动的主要因素及其作用机制；第三，通过国际比较研究，总结发达国家和新兴经济体在外汇储备管理方面的成功经验；第四，结合我国实际情况，提出优化外汇储备管理的政策建议。')
    add_body(doc, '本研究的理论意义在于：丰富和完善国际储备理论，深化对外汇储备增长与宏观经济运行关系的认识，为建立适合中国国情的外汇储备管理框架提供学理依据。实践意义在于：为政策制定者科学决策提供参考，为合理确定外汇储备适度规模提供方法支撑，为完善储备多元化投资策略提供经验借鉴。')

    add_title_h2(doc, '1.3 国内外研究现状')
    add_title_h3(doc, '1.3.1 国外研究现状')
    add_body(doc, '国际学术界对外汇储备的研究始于20世纪60年代。Heller（1966）首次从成本收益角度分析了最优外汇储备规模的决定问题，建立了基于交易动机的储备需求模型。此后，Agarwal（1971）将该模型扩展至发展中国家。Frenkel（1974）利用回归分析方法，实证检验了储备需求与进口规模、国际收支波动性等因素的关系。')
    add_body(doc, '进入21世纪，国际学者开始关注新兴市场经济体外汇储备的"过度积累"现象。Aizenman和Lee（2007）提出了"重商主义动机"假说，认为新兴经济体积累外汇储备是为了在竞争性贬值中保护出口部门。Jeanne（2007）则从"自我保险"视角分析了新兴市场国家持有高额储备的合理性。Obstfeld等（2010）进一步论证了外汇储备在预防资本突然中断方面的作用。')

    add_title_h3(doc, '1.3.2 国内研究现状')
    add_body(doc, '国内学者对外汇储备的研究主要集中在三个方面：一是适度储备规模的测算。武剑（1998）采用比例分析法，提出我国外汇储备的适度规模应为年进口额的30%-40%。王国刚（2005）运用成本收益分析框架，测算了不同政策目标下的最优储备水平。二是外汇储备增长的原因分析。余永定（2003）指出，外汇储备快速增长是"双顺差"格局的必然结果。管涛（2006）则强调汇率制度和资本管制对储备积累的重要影响。三是外汇储备管理策略。李扬（2007）提出应从投资多元化、币种结构优化等方面提高储备管理效率。')

    add_title_h2(doc, '1.4 研究内容与方法')
    add_body(doc, '本文共分五章。第一章为引言，阐述研究背景、目的意义和文献综述；第二章系统回顾我国外汇储备的发展演变历程，分析各阶段的特征和驱动因素；第三章运用协整分析和误差修正模型，实证检验外汇储备与主要宏观经济变量之间的关系；第四章通过国际比较，总结发达国家的储备管理经验及其对我国的启示；第五章为结论与展望，提出政策建议并指出研究不足。')
    add_body(doc, '本文采用的研究方法包括：（1）文献研究法，系统梳理国内外相关理论和实证研究；（2）计量分析法，运用ADF检验、Johansen协整检验和VECM模型进行实证研究；（3）比较分析法，对不同国家的储备管理模式进行横向对比；（4）规范分析与实证分析相结合，在实证检验的基础上提出规范性政策建议。')

    add_page_break(doc)

    # ══════════════════════════════════════════════
    # 六、正文：第二章
    # ══════════════════════════════════════════════
    add_title_h1(doc, '第二章 我国外汇储备的演变历程')

    add_body(doc, '改革开放以来，我国外汇储备经历了从小到大、从弱到强的发展过程。根据增长速度和驱动因素的不同，可以将其划分为以下四个阶段。')

    add_title_h2(doc, '2.1 规模较小阶段（1978—1993年）')
    add_body(doc, '1978年，我国外汇储备仅有1.67亿美元。改革开放初期，国内经济发展水平较低，出口创汇能力有限，外汇储备长期处于较低水平。1983年，我国外汇储备增至8.9亿美元。此后，由于对外汇储备规模的认识存在偏差，一度认为储备过多而加大进口力度，导致储备在1986年下降至2.07亿美元。此后逐渐恢复，到1993年底达到211.99亿美元。')
    add_body(doc, '这一阶段外汇储备规模较小的根本原因在于：我国经济总量有限，对外贸易规模较小；出口产品结构单一，以初级产品和劳动密集型产品为主；外汇管理制度高度集中，企业和居民的外汇需求受到严格限制。')

    add_title_h2(doc, '2.2 快速增长阶段（1994—1997年）')
    add_body(doc, '1994年，我国实施了以汇率并轨为核心的外汇管理体制改革，将双重汇率制度改为以市场供求为基础的、单一的、有管理的浮动汇率制度。同时，取消了外汇留成和上缴制度，实行银行结售汇制度，建立了全国统一的银行间外汇市场。这些改革措施极大地激发了出口企业的积极性，推动了外汇储备的快速增长。')
    add_body(doc, '到1997年底，我国外汇储备余额增至1398.9亿美元，较1993年的211.99亿美元增长了5.6倍，年均增长约297亿美元。这一时期外汇储备增长的主要驱动力包括：汇率并轨带来的出口竞争力提升、外商直接投资的大幅增加以及宏观经济形势总体向好。')

    add_title_h2(doc, '2.3 稳步增长阶段（1998—2000年）')
    add_body(doc, '1997年下半年，亚洲金融危机爆发，对我国出口和外商直接投资产生了一定冲击。从1998年起，我国外汇储备增量明显减缓。1998年至2000年，年增长额分别为50.97亿美元、97.15亿美元和108.99亿美元，仅相当于1997年增加额的14.6%、27.9%和31.3%。')
    add_body(doc, '尽管增长速度放缓，但我国坚持人民币不贬值的政策立场，有效维护了区域金融稳定。到2000年末，外汇储备仍增至1655.74亿美元，居世界前列。这一阶段外汇储备的稳步增长，既得益于经常项目和资本项目的"双顺差"格局，也体现了我国宏观经济政策的有效性。')

    add_title_h2(doc, '2.4 大幅增长阶段（2001年至今）')
    add_body(doc, '2001年底，我国正式加入世界贸易组织（WTO），开启了对外贸易的黄金发展期。同时，全球经济形势好转、国际产业转移加速等因素，共同推动了我国外汇储备的大幅增长。2001年至2007年，年增长额分别为465.91亿美元、742.42亿美元、1168.44亿美元、2066.81亿美元、2089亿美元、2474亿美元和4619亿美元。')
    add_body(doc, '2006年2月，我国外汇储备余额达到8536亿美元，超越日本成为全球最大的外汇储备持有国。此后，储备规模持续攀升，到2014年6月达到峰值39932亿美元。近年来，受人民币汇率波动和资本流出等因素影响，储备规模有所回落，但始终保持在3万亿美元以上。')

    add_paragraph(doc, '', '宋体', Pt(12), False, WD_ALIGN_PARAGRAPH.CENTER, 0, 0, 1.0)

    # 表格示例：中英文双语标题
    add_bilingual_caption(doc, '表2.1 我国外汇储备主要发展阶段（1978—2024年）',
                          'Table 2.1 Main development stages of China\'s foreign exchange reserves (1978-2024)')

    data = [
        ['阶段', '时间范围', '年末余额（亿美元）', '年均增长（亿美元）', '主要特征'],
        ['一', '1978—1993', '1.67—212', '14', '规模小、增速缓慢'],
        ['二', '1994—1997', '516—1399', '297', '汇率并轨、快速增长'],
        ['三', '1998—2000', '1450—1656', '86', '金融危机、稳步增长'],
        ['四', '2001至今', '2122—39932', '1735', '入世红利、大幅增长'],
    ]

    table = doc.add_table(rows=len(data), cols=5)
    table.style = 'Table Grid'
    for i, row_data in enumerate(data):
        for j, cell_text in enumerate(row_data):
            cell = table.rows[i].cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(cell_text)
            if i == 0:
                set_run_font(run, '宋体', Pt(10.5), True)
            else:
                set_run_font(run, '宋体', Pt(10.5))

    add_paragraph(doc, '注：数据来源为中国人民银行、国家外汇管理局。', '宋体', Pt(10.5),
                  False, WD_ALIGN_PARAGRAPH.LEFT, 0, 0)

    add_page_break(doc)

    # ══════════════════════════════════════════════
    # 七、正文：第三章（含图表）
    # ══════════════════════════════════════════════
    add_title_h1(doc, '第三章 外汇储备影响因素实证分析')

    add_title_h2(doc, '3.1 变量选取与数据说明')
    add_body(doc, '为定量分析影响我国外汇储备变动的主要因素，本文选取以下变量进行实证研究：（1）被解释变量——外汇储备余额（FR），单位为亿美元；（2）解释变量——进出口贸易差额（TRADE），单位为亿美元；外商直接投资（FDI），单位为亿美元；人民币兑美元汇率（EXRATE），采用间接标价法；中美利差（RATE），即中国一年期存款利率与美国联邦基金利率之差。')
    add_body(doc, '数据期间为1994年第1季度至2024年第4季度，共124个季度观测值。数据来源于中国人民银行、国家外汇管理局、国家统计局和Wind数据库。所有变量均取自然对数，以消除异方差性。')

    add_title_h2(doc, '3.2 平稳性检验')
    add_body(doc, '在进行协整分析之前，需要对各变量进行平稳性检验。本文采用ADF检验方法，检验结果如表3.1所示。')

    add_paragraph(doc, '', '宋体', Pt(12), False, WD_ALIGN_PARAGRAPH.CENTER, 0, 0, 1.0)
    add_bilingual_caption(doc, '表3.1 ADF单位根检验结果',
                          'Table 3.1 ADF unit root test results')

    adf_data = [
        ['变量', '检验形式(C,T,L)', 'ADF统计量', '临界值(5%)', '结论'],
        ['lnFR', '(C,T,2)', '-2.341', '-3.446', '不平稳'],
        ['ΔlnFR', '(C,0,1)', '-5.672', '-2.887', '平稳'],
        ['lnTRADE', '(C,T,3)', '-1.896', '-3.446', '不平稳'],
        ['ΔlnTRADE', '(C,0,2)', '-6.234', '-2.887', '平稳'],
        ['lnFDI', '(C,T,1)', '-2.156', '-3.446', '不平稳'],
        ['ΔlnFDI', '(C,0,0)', '-7.891', '-2.887', '平稳'],
        ['lnEXRATE', '(C,T,2)', '-1.543', '-3.446', '不平稳'],
        ['ΔlnEXRATE', '(C,0,1)', '-4.321', '-2.887', '平稳'],
    ]

    table = doc.add_table(rows=len(adf_data), cols=5)
    table.style = 'Table Grid'
    for i, row_data in enumerate(adf_data):
        for j, cell_text in enumerate(row_data):
            cell = table.rows[i].cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(cell_text)
            if i == 0:
                set_run_font(run, '宋体', Pt(10.5), True)
            else:
                set_run_font(run, 'Times New Roman', Pt(10.5))

    add_paragraph(doc, '注：检验形式(C,T,L)中C表示常数项，T表示趋势项，L表示滞后阶数。Δ表示一阶差分。',
                  '宋体', Pt(10.5), False, WD_ALIGN_PARAGRAPH.LEFT, 0, 0)

    add_body(doc, '由表3.1可知，所有变量的水平序列均不能拒绝存在单位根的原假设，即为非平稳序列。经一阶差分后，所有变量均在5%的显著性水平下拒绝单位根假设，表明各变量均为I(1)过程，满足协整检验的前提条件。')

    add_title_h2(doc, '3.3 协整检验与误差修正模型')
    add_body(doc, '在确认各变量均为同阶单整后，本文采用Johansen协整检验方法，检验变量之间是否存在长期均衡关系。迹检验和最大特征值检验均表明，在5%的显著性水平下，存在一个协整方程。标准化的协整方程如下：')
    add_body(doc, 'lnFR = 0.436lnTRADE + 0.287lnFDI - 1.523lnEXRATE + 0.089RATE + ε')
    add_body(doc, '协整方程表明：贸易差额每增长1%，外汇储备增加0.436%；外商直接投资每增长1%，外汇储备增加0.287%；人民币升值（汇率下降）1%，外汇储备增加1.523%；中美利差扩大1个百分点，外汇储备增加0.089%。各变量的系数符号与理论预期一致，且均具有统计显著性。')

    add_body(doc, '在协整关系的基础上，进一步建立误差修正模型（VECM），以分析短期波动特征。误差修正项的系数为-0.086，表明当外汇储备偏离长期均衡水平时，每季度约有8.6%的调整力度使其回归均衡。这一调整速度相对较慢，反映了外汇储备变动具有较强的惯性特征。')

    add_title_h2(doc, '3.4 实证结果分析')
    add_body(doc, '综合协整分析和误差修正模型的结果，可以得出以下结论：第一，贸易顺差是推动我国外汇储备增长的最主要因素，这与我国长期以来的出口导向型发展战略相吻合。第二，外商直接投资对外汇储备增长的贡献显著，反映了我国持续吸引外资的能力。第三，人民币汇率对外汇储备有显著影响，人民币升值预期会刺激资本流入，推动储备增加。第四，中美利差的影响相对较小但显著，说明利率因素在资本流动中发挥了一定作用。')

    add_body(doc, '图3.1展示了1994—2024年我国外汇储备实际值与模型拟合值的对比情况。从图中可以看出，模型能够较好地捕捉外汇储备的长期趋势，但在个别年份（如2008年金融危机期间和2015年"8·11汇改"后）存在较大偏差，说明短期冲击事件对外汇储备的影响不容忽视。')

    add_figure_placeholder(doc, '图3.1 外汇储备实际值与拟合值对比', 'Fig. 3.1 Comparison of actual and fitted values of foreign exchange reserves')

    add_page_break(doc)

    # ══════════════════════════════════════════════
    # 八、正文：第四章
    # ══════════════════════════════════════════════
    add_title_h1(doc, '第四章 国际比较与经验借鉴')

    add_body(doc, '外汇储备管理是一个全球性课题，不同国家根据自身经济体制和发展阶段，形成了各具特色的管理模式。本章选取日本、新加坡和挪威三个具有代表性的国家，分析其外汇储备管理的实践经验及其对我国的启示。')

    add_title_h2(doc, '4.1 日本外汇储备管理')
    add_body(doc, '日本是全球第二大外汇储备持有国。其储备管理以财务省为主导，日本银行（央行）负责具体操作。日本采取"三层次"管理模式：第一层为流动性资产，用于干预外汇市场；第二层为投资性资产，追求中长期收益；第三层为战略性资产，服务于国家长期利益。日本的经验表明，合理的储备分层管理有助于兼顾流动性需求和收益性目标。')

    add_title_h2(doc, '4.2 新加坡外汇储备管理')
    add_body(doc, '新加坡的外汇储备由新加坡金融管理局（MAS）和政府投资公司（GIC）分别管理。MAS主要管理用于货币政策操作的储备部分，强调流动性和安全性；GIC则负责管理长期投资部分，投资范围涵盖全球股票、债券和另类资产。新加坡模式的核心特点是将储备管理与主权财富基金有机结合，在确保储备安全性的前提下追求更高的长期回报。')

    add_title_h2(doc, '4.3 挪威外汇储备管理')
    add_body(doc, '挪威通过挪威央行投资管理局（NBIM）管理其主权财富基金——全球政府养老基金（GPFG），规模超过1.5万亿美元，是全球最大的主权财富基金之一。NBIM采用被动投资与主动管理相结合的策略，投资于全球70多个国家的股票、债券和房地产。挪威模式的最大亮点在于建立了严格的治理框架和透明的信息披露机制，确保基金管理的规范性和公信力。')

    add_title_h2(doc, '4.4 经验总结与启示')
    add_body(doc, '通过对上述三国储备管理经验的比较分析，可以总结出以下启示：第一，建立多层次、专业化的管理架构，实现流动性管理与投资管理的有机分离。第二，完善治理结构和风险控制机制，建立独立的投资决策和监督体系。第三，推进投资多元化，在币种、资产类别和地域上适度分散风险。第四，提高管理透明度，建立规范的信息披露制度。第五，将外汇储备管理与国家发展战略有机结合，发挥储备资产的综合效益。')

    add_paragraph(doc, '', '宋体', Pt(12), False, WD_ALIGN_PARAGRAPH.CENTER, 0, 0, 1.0)
    add_bilingual_caption(doc, '表4.1 三国外汇储备管理模式比较',
                          'Table 4.1 Comparison of foreign exchange reserve management models of three countries')

    comp_data = [
        ['项目', '日本', '新加坡', '挪威'],
        ['管理机构', '财务省+日本银行', 'MAS+GIC', 'NBIM'],
        ['管理目标', '流动性+收益性', '安全性+长期回报', '长期价值增长'],
        ['投资范围', '主权债券为主', '全球多资产', '全球多资产'],
        ['透明度', '中等', '较高', '很高'],
        ['特色', '市场干预导向', '主权财富基金', '养老金模式'],
    ]

    table = doc.add_table(rows=len(comp_data), cols=4)
    table.style = 'Table Grid'
    for i, row_data in enumerate(comp_data):
        for j, cell_text in enumerate(row_data):
            cell = table.rows[i].cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(cell_text)
            if i == 0:
                set_run_font(run, '宋体', Pt(10.5), True)
            else:
                set_run_font(run, '宋体', Pt(10.5))

    add_page_break(doc)

    # ══════════════════════════════════════════════
    # 九、正文：第五章 结论与展望
    # ══════════════════════════════════════════════
    add_title_h1(doc, '第五章 结论与展望')

    add_title_h2(doc, '5.1 主要结论')
    add_body(doc, '本文围绕我国外汇储备的增长机理、影响因素及管理策略展开了系统研究，主要得出以下结论：')
    add_body(doc, '第一，我国外汇储备的增长具有明显的阶段性特征，贸易顺差和外商直接投资是最主要的驱动因素。从1978年的不足2亿美元到2014年近4万亿美元的峰值，外汇储备的快速增长反映了我国经济实力的增强和国际地位的提升。')
    add_body(doc, '第二，协整分析表明，外汇储备与贸易差额、外商直接投资之间存在显著的正向长期均衡关系，人民币汇率和中美利差也对储备变动具有重要影响。短期内，储备向均衡水平调整的速度约为每季度8.6%，具有较强的惯性。')
    add_body(doc, '第三，国际比较表明，成熟的外汇储备管理体系应具备多层次管理架构、专业化投资团队、严格的风险控制机制和较高的透明度。日本、新加坡和挪威的经验为我国提供了有益参考。')

    add_title_h2(doc, '5.2 政策建议')
    add_body(doc, '基于以上研究结论，本文提出以下政策建议：')
    add_body(doc, '第一，优化储备资产结构，推进投资多元化。在确保安全性和流动性的前提下，适当增加黄金、大宗商品等实物资产的配置比例，拓展投资的地域和行业范围，降低单一资产类别的集中风险。')
    add_body(doc, '第二，增强人民币汇率弹性，发挥汇率调节作用。进一步完善人民币汇率形成机制，扩大汇率浮动区间，减少央行在外汇市场的常态化干预，让汇率成为调节国际收支的自动稳定器。')
    add_body(doc, '第三，稳步推进人民币国际化，从源头上减少储备积累压力。扩大人民币在贸易结算和投资中的使用，推动人民币加入特别提款权（SDR）货币篮子的实际应用，降低对美元的过度依赖。')
    add_body(doc, '第四，完善储备管理的制度建设。借鉴挪威等国经验，建立独立的外汇储备投资管理机构，完善内部治理结构和风险控制框架，提高管理的专业化水平和透明度。')

    add_title_h2(doc, '5.3 研究不足与展望')
    add_body(doc, '本文仍存在以下不足：第一，在数据方面，受季度数据可得性限制，部分变量的样本期较短，可能影响实证结果的稳健性。第二，在方法方面，本文主要采用线性协整模型，未充分考虑变量之间的非线性关系和结构突变问题。第三，在国际比较方面，受篇幅所限，仅选取了三个国家进行分析，案例的代表性有待进一步增强。')
    add_body(doc, '未来研究可从以下方向深入：一是采用非线性协整和门限模型，检验储备变动的非对称特征；二是利用高频数据分析短期资本流动对储备变动的冲击效应；三是将数字货币发展、国际货币体系变革等新因素纳入分析框架；四是建立动态随机一般均衡（DSGE）模型，从一般均衡视角评估不同储备管理策略的宏观经济效应。')

    add_page_break(doc)

    # ══════════════════════════════════════════════
    # 十、参考文献
    # ══════════════════════════════════════════════
    add_paragraph(doc, '参 考 文 献', '宋体', Pt(14), True, WD_ALIGN_PARAGRAPH.CENTER, 0, 6)

    refs = [
        '[1] Heller H R. Optimal international reserves[J]. The Economic Journal, 1966, 76(302): 296-311.',
        '[2] Agarwal J P. Optimal monetary reserves for developing countries[J]. Review of World Economics, 1971, 107(1): 76-91.',
        '[3] Frenkel J A. The demand for international reserves by developed and less-developed countries[J]. Economica, 1974, 41(161): 14-24.',
        '[4] Aizenman J, Lee J. International reserves: precautionary versus mercantilist views, theory and evidence[J]. Open Economies Review, 2007, 18(2): 191-214.',
        '[5] Jeanne O. International reserves in emerging market countries: too much of a good thing?[J]. Brookings Papers on Economic Activity, 2007(1): 1-55.',
        '[6] Obstfeld M, Shambaugh J C, Taylor A M. Financial stability, the trilemma, and international reserves[J]. American Economic Journal: Macroeconomics, 2010, 2(2): 57-94.',
        '[7] 余永定. 关于外汇储备和国际收支的几个问题[J]. 国际金融研究, 2003(10): 2-8.',
        '[8] 武剑. 我国外汇储备规模的分析与测算[J]. 经济研究, 1998(6): 34-40.',
        '[9] 王国刚. 中国外汇储备的经济分析[J]. 经济研究, 2005(3): 18-29.',
        '[10] 管涛. 中国外汇储备管理体制改革的路径选择[J]. 国际金融研究, 2006(9): 4-10.',
        '[11] 李扬. 中国外汇储备管理的战略思考[J]. 金融研究, 2007(3): 1-13.',
        '[12] 刘国钧, 郑如斯. 中国书的故事[M]. 北京: 中国青年出版社, 1979: 115.',
        '[13] 张和生. 地质力学系统理论[D]. 太原: 太原理工大学, 1998.',
        '[14] 冯西桥. 核反应堆压力容器的LBB分析[R]. 北京: 清华大学核能技术设计研究院, 1997.',
        '[15] GB/T 16159-1996. 汉语拼音正词法基本规则[S]. 北京: 中国标准出版社, 1996.',
    ]
    for ref in refs:
        add_ref(doc, ref)

    add_page_break(doc)

    # ══════════════════════════════════════════════
    # 十一、致谢
    # ══════════════════════════════════════════════
    add_paragraph(doc, '致  谢', '宋体', Pt(18), True, WD_ALIGN_PARAGRAPH.CENTER, 0, 6)
    add_body(doc, '时光荏苒，三年的研究生生涯即将画上句号。回首这段求学岁月，心中感慨万千。从论文选题的迷茫到研究思路的逐步清晰，从数据收集的艰辛到模型构建的反复调试，每一步都凝聚着师长的指导、同学的帮助和家人的支持。')
    add_body(doc, '首先，我要衷心感谢我的导师XXX教授。从论文的选题、开题、撰写到修改，先生始终给予我耐心细致的指导。先生严谨的治学态度、渊博的学识和高尚的人格魅力，不仅为我的学术研究指明了方向，更将成为我未来人生道路上的宝贵精神财富。在此，谨向先生致以最崇高的敬意和最诚挚的感谢！')
    add_body(doc, '其次，我要感谢经济学院的各位老师。XXX老师、XXX老师在课程教学中为我打下了扎实的理论基础，XXX老师在研究方法上给予了我悉心的指导。感谢各位评审专家在百忙之中审阅本文并提出宝贵意见。')
    add_body(doc, '感谢同门的师兄师姐和师弟师妹们。在研究过程中，我们相互交流、共同进步，那些在实验室里共同奋斗的日日夜夜，将成为我最美好的回忆。特别感谢XXX同学在数据处理方面给予的帮助，感谢XXX同学在论文修改中提出的建议。')
    add_body(doc, '最后，我要感谢我的父母和家人。感谢他们多年来的默默付出和无私支持，是他们的理解与鼓励，让我能够安心学业、潜心研究。')
    add_body(doc, '纸短情长，无法一一列举所有帮助过我的人。在论文完成之际，谨向所有关心和帮助过我的老师、同学和朋友们，致以最衷心的感谢！')

    add_page_break(doc)

    # ══════════════════════════════════════════════
    # 十二、附录：攻读学位期间发表的论文
    # ══════════════════════════════════════════════
    add_paragraph(doc, '附录A 攻读学位期间发表的论文', '宋体', Pt(18), True, WD_ALIGN_PARAGRAPH.CENTER, 0, 12)
    add_paragraph(doc, '', '宋体', Pt(12), False, WD_ALIGN_PARAGRAPH.CENTER, 0, 0, 1.0)

    papers = [
        '[1] XXX（第一作者）. 基于协整分析的我国外汇储备增长影响因素研究[J]. 湖南科技大学学报（社会科学版）, 2025, 28(3): 45-53.',
        '[2] XXX（第一作者）. 新兴市场国家外汇储备管理的国际比较与启示[J]. 金融理论与实践, 2025(6): 78-86.',
        '[3] XXX（第二作者，导师第一）. 经济全球化背景下我国外汇储备适度规模研究[J]. 经济问题探索, 2024(12): 112-120.',
    ]
    for paper in papers:
        add_ref(doc, paper)

    # ── 保存 ──
    doc.save(filepath)
    print(f"  [OK] 硕士论文示例文档已保存到: {filepath}")
    return filepath


def main():
    outpath = sys.argv[1] if len(sys.argv) > 1 else '硕士论文格式示例.docx'
    # 确保输出到 examples 目录
    if not os.path.dirname(outpath):
        outpath = os.path.join(os.path.dirname(__file__), '..', 'examples', outpath)
    os.makedirs(os.path.dirname(os.path.abspath(outpath)), exist_ok=True)
    generate(outpath)


if __name__ == '__main__':
    main()
