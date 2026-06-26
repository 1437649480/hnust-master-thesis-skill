# -*- coding: utf-8 -*-
"""
湖南科技大学硕士学位论文格式检查脚本
基于《湖南科技大学研究生学位论文撰写规范》
用法: python check_thesis.py <docx文件路径>
"""

import sys, re, os
from docx import Document
from docx.shared import Pt, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# ─── 常量定义（硕士论文规范） ───
SPEC = {
    # 1. 页面设置：A4纸(210mm×297mm)，页眉页脚距边界20mm，上下30mm，左右25mm
    'page': {'width': Cm(21), 'height': Cm(29.7), 'top': Cm(3), 'bottom': Cm(3),
             'left': Cm(2.5), 'right': Cm(2.5), 'header': Cm(2), 'footer': Cm(2)},
    # 2. 页眉：五号宋体(10.5pt)居中
    'header_toc': {'text': '湖南科技大学硕士学位论文', 'font': '宋体', 'size': Pt(10.5),
                   'align': WD_ALIGN_PARAGRAPH.CENTER},
    'header_odd': {'text': '湖南科技大学硕士学位论文', 'font': '宋体', 'size': Pt(10.5),
                   'align': WD_ALIGN_PARAGRAPH.CENTER},
    # 3. 页码：5号Times New Roman居中，格式"-X-"
    'page_num': {'font': 'Times New Roman', 'size': Pt(10.5), 'align': WD_ALIGN_PARAGRAPH.CENTER},
    # 4. 字体字号
    'h1': {'font': '宋体', 'size': Pt(18), 'bold': True, 'align': WD_ALIGN_PARAGRAPH.CENTER,
           'before': Pt(12), 'after': Pt(12)},       # 章标题：小2号宋体加粗，上下各空一行
    'h2': {'font': '宋体', 'size': Pt(15), 'bold': True, 'align': None,
           'before': Pt(6), 'after': Pt(6)},          # 节标题：小3号宋体加粗，上下各空0.5行
    'h3': {'font': '宋体', 'size': Pt(14), 'bold': True, 'align': None,
           'before': Pt(6), 'after': Pt(6)},          # 二级标题：4号宋体加粗，上下各空0.5行
    'h4': {'font': '宋体', 'size': Pt(12), 'bold': True, 'align': None,
           'before': Pt(6), 'after': Pt(6)},          # 三级标题：小4号宋体加粗，上下各空0.5行
    'h5': {'font': '宋体', 'size': Pt(12), 'bold': True, 'align': None,
           'before': Pt(0), 'after': Pt(0)},          # 条款项：小4号宋体加粗，上下不空
    'body': {'font': '宋体', 'size': Pt(12), 'bold': False, 'align': None,
             'line_spacing': 1.25, 'first_line_indent': Pt(24)},  # 正文：小4号宋体，1.25倍行距
    # 摘要
    'abstract_title': {'font': '宋体', 'size': Pt(18), 'bold': True, 'align': WD_ALIGN_PARAGRAPH.CENTER},
    'abstract_body': {'font': '宋体', 'size': Pt(12), 'bold': False},
    'abstract_kw_label': {'font': '黑体', 'size': Pt(12), 'bold': False},  # 关键词标签：小4号黑体
    'abstract_kw': {'font': '宋体', 'size': Pt(12), 'bold': False},
    'en_abstract_title': {'font': 'Times New Roman', 'size': Pt(18), 'bold': True,
                          'align': WD_ALIGN_PARAGRAPH.CENTER},
    'en_abstract_body': {'font': 'Times New Roman', 'size': Pt(12), 'bold': False},
    'en_kw_label': {'font': 'Times New Roman', 'size': Pt(12), 'bold': True},
    'en_kw': {'font': 'Times New Roman', 'size': Pt(12), 'bold': False},
    # 目录标题
    'toc_title': {'font': '黑体', 'size': Pt(16), 'bold': False},
    # 6. 参考文献
    'ref_title': {'font': '宋体', 'size': Pt(12), 'bold': True, 'align': WD_ALIGN_PARAGRAPH.CENTER},
    'ref_entry': {'font': '宋体', 'size': Pt(9), 'bold': False},        # 小5号宋体(9pt)
    'ref_entry_en': {'font': 'Times New Roman', 'size': Pt(9), 'bold': False},
    # 其他
    'cover': {'font': '宋体', 'size': Pt(22), 'bold': True},
    'thanks_title': {'font': '宋体', 'size': Pt(18), 'bold': True, 'align': WD_ALIGN_PARAGRAPH.CENTER},
    'appendix_title': {'font': '宋体', 'size': Pt(18), 'bold': True, 'align': WD_ALIGN_PARAGRAPH.CENTER},
    'table_cell': {'font_cn': '宋体', 'font_en': 'Times New Roman', 'size': Pt(10.5)},
}

TOLERANCE = 50000  # EMU 容差


# ─── 工具函数 ───
def get_text(para):
    """获取段落纯文本"""
    return para.text.strip()


def get_run_info(para):
    """获取段落中第一个有文本的run"""
    for run in para.runs:
        if run.text.strip():
            return run
    return para.runs[0] if para.runs else None


def match_h1(text):
    """匹配章标题: 第X章"""
    return bool(re.match(r'^第[一二三四五六七八九十\d]+章\s', text))


def match_h2(text):
    """匹配节标题: X.Y"""
    return bool(re.match(r'^[\d]+\.[\d]+\s', text)) and not match_h3(text)


def match_h3(text):
    """匹配二级标题: X.Y.Z"""
    return bool(re.match(r'^[\d]+\.[\d]+\.[\d]+\s', text)) and not match_h4(text)


def match_h4(text):
    """匹配三级标题: X.Y.Z.W"""
    return bool(re.match(r'^[\d]+\.[\d]+\.[\d]+\.[\d]+\s', text))


def match_h5(text):
    """匹配款项: (1) 或 ①"""
    return bool(re.match(r'^[（(][\d一二三四五六七八九十]+[）)]', text))


def match_ref(text):
    """匹配参考文献条目: [1]"""
    return bool(re.match(r'^\[\d+\]', text))


def match_appendix(text):
    """匹配附录标题"""
    return bool(re.match(r'^附录[A-Z]', text))


def match_thanks_title(text):
    """匹配致谢标题"""
    return text in ('致 谢', '致谢') or bool(re.match(r'^致\s+谢$', text))


def emu_to_cm(v):
    return v / 360000


def check_val(actual, expected, tolerance=TOLERANCE):
    if isinstance(expected, (int, float)) and isinstance(actual, (int, float)):
        return abs(actual - expected) <= tolerance
    return actual == expected


# ─── 检查器类 ───
class ThesisChecker:
    def __init__(self, filepath):
        self.filepath = filepath
        self.doc = Document(filepath)
        self.errors = []
        self.warnings = []
        self.passed = []
        self.sections = self.doc.sections
        self.paragraphs = self.doc.paragraphs
        self.toc_start = None
        self.toc_end = None
        self.cover_end = None
        self._scan_regions()

    def check(self, condition, category, detail):
        if condition:
            self.passed.append(f"[{category}] {detail}")
        else:
            self.errors.append(f"[{category}] {detail}")

    def check_warn(self, condition, category, detail):
        if condition:
            self.passed.append(f"[{category}] {detail}")
        else:
            self.warnings.append(f"[{category}] {detail}")

    def _scan_regions(self):
        """预扫描文档区域：识别目录、扉页等边界"""
        all_text = [get_text(p) for p in self.paragraphs]
        for i, t in enumerate(all_text):
            if t in ('目 录', '目录') or re.match(r'^目\s+录$', t):
                self.toc_start = i
                break
        if self.toc_start is not None:
            for i in range(self.toc_start + 1, len(all_text)):
                text = all_text[i]
                if match_h1(text) and '\t' not in text:
                    self.toc_end = i
                    break
        for i, t in enumerate(all_text):
            if t in ('摘  要', '摘 要', '摘要') or re.match(r'^摘\s+要$', t) or t.strip().upper() == 'ABSTRACT':
                self.cover_end = i
                break

    def _is_in_toc(self, para_idx):
        if self.toc_start is None or self.toc_end is None:
            return False
        return self.toc_start <= para_idx < self.toc_end

    def _is_in_cover(self, para_idx):
        if self.cover_end is None:
            return False
        return para_idx < self.cover_end

    # ── 1. 页面设置 ──
    def check_page_setup(self):
        sec = self.sections[0]
        self.check(abs(sec.page_width - SPEC['page']['width']) <= TOLERANCE,
                   '页面', f'A4宽度 210mm (实际: {emu_to_cm(sec.page_width):.1f}mm)')
        self.check(abs(sec.page_height - SPEC['page']['height']) <= TOLERANCE,
                   '页面', f'A4高度 297mm (实际: {emu_to_cm(sec.page_height):.1f}mm)')
        self.check(abs(sec.top_margin - SPEC['page']['top']) <= TOLERANCE,
                   '页面', f'上页边距 30mm (实际: {emu_to_cm(sec.top_margin):.1f}mm)')
        self.check(abs(sec.bottom_margin - SPEC['page']['bottom']) <= TOLERANCE,
                   '页面', f'下页边距 30mm (实际: {emu_to_cm(sec.bottom_margin):.1f}mm)')
        self.check(abs(sec.left_margin - SPEC['page']['left']) <= TOLERANCE,
                   '页面', f'左页边距 25mm (实际: {emu_to_cm(sec.left_margin):.1f}mm)')
        self.check(abs(sec.right_margin - SPEC['page']['right']) <= TOLERANCE,
                   '页面', f'右页边距 25mm (实际: {emu_to_cm(sec.right_margin):.1f}mm)')
        self.check(abs(sec.header_distance - SPEC['page']['header']) <= TOLERANCE,
                   '页面', f'页眉距边界 20mm (实际: {emu_to_cm(sec.header_distance):.1f}mm)')
        self.check(abs(sec.footer_distance - SPEC['page']['footer']) <= TOLERANCE,
                   '页面', f'页脚距边界 20mm (实际: {emu_to_cm(sec.footer_distance):.1f}mm)')

    # ── 2. 页眉 ──
    def check_header(self):
        for si, sec in enumerate(self.sections):
            header = sec.header
            if header.is_linked_to_previous:
                continue  # 跳过链接到前一节的页眉
            paras = header.paragraphs
            if not paras or not paras[0].text.strip():
                self.check(False, f'页眉(节{si+1})', '页眉为空')
                continue
            p = paras[0]
            text = p.text.strip()
            # 硕士论文页眉均为"湖南科技大学硕士学位论文"
            self.check(SPEC['header_odd']['text'] in text,
                       f'页眉(节{si+1})', f'页眉内容应为"{SPEC["header_odd"]["text"]}" (实际: "{text}")')
            self.check(p.alignment == SPEC['header_odd']['align'],
                       f'页眉(节{si+1})', '页眉应居中')
            run = get_run_info(p)
            if run:
                self.check(run.font.size and abs(run.font.size - SPEC['header_odd']['size']) <= TOLERANCE,
                           f'页眉(节{si+1})', f'字号应为五号/10.5pt (实际: {run.font.size})')
                self.check(run.font.name == SPEC['header_odd']['font'],
                           f'页眉(节{si+1})', f'字体应为宋体 (实际: {run.font.name})')

    # ── 3. 页码 ──
    def check_page_number(self):
        for si, sec in enumerate(self.sections):
            footer = sec.footer
            sectPr = sec._sectPr
            pgnt = sectPr.find(qn('w:pgNumType'))
            if pgnt is not None:
                fmt = pgnt.get(qn('w:fmt'))
                if si == 0 and len(self.sections) > 1:
                    self.check(fmt == 'lowerRoman',
                               f'页码(节{si+1})', f'前置部分应为罗马数字 (实际: {fmt})')
                else:
                    self.check(fmt == 'decimal',
                               f'页码(节{si+1})', f'正文部分应为阿拉伯数字 (实际: {fmt})')
            if not footer.is_linked_to_previous:
                paras = footer.paragraphs
                if paras:
                    p = paras[0]
                    self.check(p.alignment == SPEC['page_num']['align'],
                               f'页码(节{si+1})', '页码应居中')
                    run = get_run_info(p)
                    if run:
                        self.check(run.font.size and abs(run.font.size - SPEC['page_num']['size']) <= TOLERANCE,
                                   f'页码(节{si+1})', f'页码字号应为5号/10.5pt (实际: {run.font.size})')
                        rFonts = run._element.find(qn('w:rPr'))
                        if rFonts is not None:
                            rFonts_elem = rFonts.find(qn('w:rFonts'))
                            if rFonts_elem is not None:
                                ascii_f = rFonts_elem.get(qn('w:ascii')) or run.font.name
                                self.check(ascii_f == 'Times New Roman',
                                           f'页码(节{si+1})', f'页码字体应为Times New Roman (实际: {ascii_f})')

    # ── 4. 字体字号段落格式 ──
    def check_font_paragraph(self):
        for i, para in enumerate(self.paragraphs):
            text = get_text(para)
            if not text:
                continue
            if self._is_in_toc(i):
                continue
            if self._is_in_cover(i):
                continue

            # 章标题：小2号宋体加粗(18pt)
            if match_h1(text):
                self._check_heading(para, 'h1', f'章标题: {text[:30]}')
                self.check(para.paragraph_format.alignment == SPEC['h1']['align'],
                           '格式', f'章标题"{text[:20]}"应居中')
                self._check_spacing(para, 'h1', f'章标题"{text[:20]}"')
                continue

            # 节标题：小3号宋体加粗(15pt)
            if match_h2(text):
                self._check_heading(para, 'h2', f'节标题: {text[:30]}')
                self._check_spacing(para, 'h2', f'节标题"{text[:20]}"')
                continue

            # 二级标题：4号宋体加粗(14pt)
            if match_h3(text):
                self._check_heading(para, 'h3', f'二级标题: {text[:30]}')
                self._check_spacing(para, 'h3', f'二级标题"{text[:20]}"')
                continue

            # 三级标题：小4号宋体加粗(12pt)
            if match_h4(text):
                self._check_heading(para, 'h4', f'三级标题: {text[:30]}')
                self._check_spacing(para, 'h4', f'三级标题"{text[:20]}"')
                continue

            # 条款项：小4号宋体加粗(12pt)，上下不空
            if match_h5(text):
                self._check_heading(para, 'h5', f'条款项: {text[:30]}')
                self.check(para.paragraph_format.first_line_indent is not None,
                           '格式', f'条款项"{text[:20]}"应有首行缩进')
                continue

            # 参考文献标题：小4号加粗居中
            if text in ('参考文献', '参 考 文 献'):
                self._check_heading(para, 'ref_title', '参考文献标题')
                self.check(para.paragraph_format.alignment == SPEC['ref_title']['align'],
                           '格式', '参考文献标题应居中')
                continue

            # 参考文献条目：小5号宋体(9pt)
            if match_ref(text):
                self._check_ref_entry(para, f'参考文献: {text[:40]}')
                self.check(para.paragraph_format.first_line_indent is None or
                           para.paragraph_format.first_line_indent == 0,
                           '格式', f'参考文献"{text[:30]}"应顶格')
                self._check_ref_entry_mixed_font(para, text)
                continue

            # 致谢标题
            if match_thanks_title(text):
                self._check_heading(para, 'thanks_title', '致谢标题')
                self.check(para.paragraph_format.alignment == SPEC['thanks_title']['align'],
                           '格式', '致谢标题应居中')
                continue

            # 附录标题
            if match_appendix(text):
                self.check(para.paragraph_format.alignment == SPEC['appendix_title']['align'],
                           '格式', f'附录标题"{text[:20]}"应居中')
                run = get_run_info(para)
                if run:
                    self.check(run.font.size and abs(run.font.size - SPEC['appendix_title']['size']) <= TOLERANCE,
                               '字号', f'附录标题"{text[:20]}"字号应为小二号/18pt')
                    self.check(run.bold, '加粗', f'附录标题"{text[:20]}"应加粗')
                continue

            # 摘要标题
            if re.match(r'^摘\s*要$', text):
                self._check_heading(para, 'abstract_title', '中文摘要标题')
                self.check(para.paragraph_format.alignment == SPEC['abstract_title']['align'],
                           '格式', '摘要标题应居中')
                continue

            # 英文摘要标题
            if text.strip().upper() == 'ABSTRACT':
                self._check_heading(para, 'en_abstract_title', '英文摘要标题')
                self.check(para.paragraph_format.alignment == SPEC['en_abstract_title']['align'],
                           '格式', 'ABSTRACT应居中')
                continue

            # 关键词：小4号黑体字
            if text.startswith('关键词') or text.startswith('关键词：'):
                self._check_kw_format(para)
                continue
            if text.startswith('Keywords') or text.startswith('KeyWords'):
                self._check_en_kw_format(para)
                continue

            # 正文
            self._check_body_format(para, text)

    def _check_heading(self, para, level, label):
        spec = SPEC[level]
        run = get_run_info(para)
        if not run:
            return
        # 参考文献标题和条目特殊处理
        if level == 'ref_entry':
            self._check_ref_entry(para, label)
            return
        self.check(run.font.name and run.font.name == spec['font'],
                   '字体', f'{label} 字体应为{spec["font"]} (实际: {run.font.name})')
        self.check(run.font.size and abs(run.font.size - spec['size']) <= TOLERANCE,
                   '字号', f'{label} 字号应为{spec["size"]/12700:.0f}pt (实际: {run.font.size/12700 if run.font.size else "None"}pt)')
        if spec.get('bold'):
            self.check(run.bold, '加粗', f'{label} 应为加粗')

    def _check_ref_entry(self, para, label):
        """检查参考文献条目字体：小5号宋体(9pt)"""
        spec = SPEC['ref_entry']
        run = get_run_info(para)
        if not run:
            return
        rPr = run._element.find(qn('w:rPr'))
        ascii_font = run.font.name
        ea_font = run.font.name
        if rPr is not None:
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is not None:
                ascii_font = rFonts.get(qn('w:ascii')) or ascii_font
                ea_font = rFonts.get(qn('w:eastAsia')) or ea_font
        self.check(ascii_font == 'Times New Roman' or ascii_font == '宋体',
                   '字体', f'{label} 西文字体应为Times New Roman (实际: {ascii_font})')
        self.check(ea_font == '宋体',
                   '字体', f'{label} 中文字体应为宋体 (实际: {ea_font})')
        self.check(run.font.size and abs(run.font.size - spec['size']) <= TOLERANCE,
                   '字号', f'{label} 字号应为小5号/9pt (实际: {run.font.size/12700 if run.font.size else "None"}pt)')

    def _check_spacing(self, para, level, label):
        """检查段前段后间距"""
        spec = SPEC[level]
        pf = para.paragraph_format
        before_ok = pf.space_before is None or abs(pf.space_before - spec['before']) <= TOLERANCE
        before_pt = int(pf.space_before / 12700) if pf.space_before else 0
        after_pt = int(pf.space_after / 12700) if pf.space_after else 0
        self.check(before_ok, '段间距', f'{label} 段前应{int(spec["before"]/12700)}pt (实际: {before_pt}pt)')
        after_ok = pf.space_after is None or abs(pf.space_after - spec['after']) <= TOLERANCE
        self.check(after_ok, '段间距', f'{label} 段后应{int(spec["after"]/12700)}pt (实际: {after_pt}pt)')

    def _check_body_format(self, para, text):
        """检查正文格式：小4号宋体(12pt)，1.25倍行距"""
        if re.match(r'^(表|图)[\d]+\.[\d]+', text):
            return
        if text.startswith('资料来源') or text.startswith('单位：'):
            return
        # 字体字号
        run = get_run_info(para)
        if run:
            if run.font.size:
                self.check(abs(run.font.size - SPEC['body']['size']) <= TOLERANCE,
                           '字号', f'正文"{text[:20]}"字号应为小4号/12pt (实际: {run.font.size/12700 if run.font.size else "None"}pt)')
        # 行距
        pf = para.paragraph_format
        line_spacing = pf.line_spacing
        if line_spacing is not None:
            self.check(abs(line_spacing - SPEC['body']['line_spacing']) <= 0.05,
                       '行距', f'正文"{text[:20]}"行距应为1.25倍 (实际: {line_spacing})')
        # 首行缩进
        if pf.first_line_indent is not None:
            self.check(abs(pf.first_line_indent - SPEC['body']['first_line_indent']) <= TOLERANCE,
                       '缩进', f'正文"{text[:20]}"首行缩进应为2字符/24pt')
        # 段前段后
        if pf.space_before is not None and pf.space_before > 0:
            before_pt = int(pf.space_before / 12700)
            self.check(False, '段间距', f'正文"{text[:20]}"段前不应有空行 (实际: {before_pt}pt)')
        if pf.space_after is not None and pf.space_after > 0:
            after_pt = int(pf.space_after / 12700)
            self.check(False, '段间距', f'正文"{text[:20]}"段后不应有空行 (实际: {after_pt}pt)')

    def _check_kw_format(self, para):
        """检查中文关键词：3-8个，小4号黑体字"""
        text = get_text(para)
        for run in para.runs:
            if '关键词' in run.text:
                self.check(run.font.name == '黑体',
                           '字体', f'关键词标签应黑体 (实际: {run.font.name})')
                self.check(run.font.size and abs(run.font.size - Pt(12)) <= TOLERANCE,
                           '字号', '关键词标签应小4号/12pt')

    def _check_en_kw_format(self, para):
        """检查英文关键词格式"""
        for run in para.runs:
            if 'Keyword' in run.text or 'KEYWORD' in run.text.upper():
                self.check(run.font.name == 'Times New Roman',
                           '字体', f'Keywords标签应Times New Roman (实际: {run.font.name})')
                self.check(run.bold, '加粗', 'Keywords标签应加粗')

    # ── 5. 关键词数量（3-8个） ──
    def check_keyword_count(self):
        for para in self.paragraphs:
            text = get_text(para)
            if text.startswith('关键词') or text.startswith('关键词：'):
                kw_text = re.sub(r'^关键词[：:]?\s*', '', text)
                keywords = [k.strip() for k in re.split(r'[；;]', kw_text) if k.strip()]
                self.check(3 <= len(keywords) <= 8,
                           '关键词', f'中文关键词应为3~8个 (实际: {len(keywords)}个)')
            if text.lower().startswith('keywords'):
                kw_text = re.sub(r'^keywords[：:]\s*', '', text, flags=re.IGNORECASE)
                keywords = [k.strip() for k in re.split(r'[；;]', kw_text) if k.strip()]
                self.check(3 <= len(keywords) <= 8,
                           '关键词', f'英文关键词应为3~8个 (实际: {len(keywords)}个)')

    # ── 7. 图表标题需中英文双语 ──
    def check_bilingual_caption(self):
        """检查图表标题是否中英文双语"""
        for para in self.paragraphs:
            text = get_text(para)
            if not text:
                continue
            # 匹配图标题: 图X.Y
            if re.match(r'^图\s*[\d]+\.[\d]+', text):
                has_cn = bool(re.search(r'[\u4e00-\u9fff]', text))
                has_en = bool(re.search(r'[a-zA-Z]{2,}', text))
                self.check_warn(has_cn and has_en,
                                '插图', f'图标题应中英文双语: "{text[:40]}"')
            # 匹配表标题: 表X.Y
            if re.match(r'^表\s*[\d]+\.[\d]+', text):
                has_cn = bool(re.search(r'[\u4e00-\u9fff]', text))
                has_en = bool(re.search(r'[a-zA-Z]{2,}', text))
                self.check_warn(has_cn and has_en,
                                '表格', f'表标题应中英文双语: "{text[:40]}"')

    # ── 8. 论文题目≤25字 ──
    def check_paper_title_length(self):
        """检查论文题目长度"""
        for i, para in enumerate(self.paragraphs[:30]):
            text = get_text(para)
            if match_h1(text):
                title_only = re.sub(r'^第[一二三四五六七八九十\d]+章\s*', '', text)
                self.check_warn(len(title_only) <= 25,
                                '标题', f'章标题"{text[:25]}"应≤25字 (实际: {len(title_only)}字)')
                self.check_warn(not re.search(r'[，。；：？！、]', title_only),
                                '标题', f'章标题"{text[:20]}"不应使用标点符号')

    # ── 9. 摘要中文600字左右 ──
    def check_abstract_word_count(self):
        """检查中文摘要字数（约600字）"""
        in_cn_abstract = False
        abstract_text = []
        for para in self.paragraphs:
            text = get_text(para)
            if re.match(r'^摘\s*要$', text):
                in_cn_abstract = True
                continue
            if text.startswith('关键词') or text.startswith('Keywords'):
                in_cn_abstract = False
                continue
            if text.strip().upper() == 'ABSTRACT':
                in_cn_abstract = False
                continue
            if in_cn_abstract and text:
                abstract_text.append(text)
        if abstract_text:
            full_text = ''.join(abstract_text)
            cn_chars = len(re.findall(r'[\u4e00-\u9fff]', full_text))
            self.check_warn(500 <= cn_chars <= 800,
                            '摘要', f'中文摘要约600字 (实际: {cn_chars}字)')

    # ── 10. 外文参考文献不少于总文献数的1/3 ──
    def check_foreign_ref_ratio(self):
        """检查外文参考文献占比"""
        in_refs = False
        total = 0
        foreign = 0
        for para in self.paragraphs:
            text = get_text(para)
            if not text:
                continue
            if text in ('参考文献', '参 考 文 献'):
                in_refs = True
                continue
            if match_thanks_title(text) or match_appendix(text):
                in_refs = False
                continue
            if in_refs and match_ref(text):
                total += 1
                # 判断是否为外文文献：不含中文字符或以英文作者名开头
                has_cn = bool(re.search(r'[\u4e00-\u9fff]', text))
                if not has_cn:
                    foreign += 1
        if total > 0:
            ratio = foreign / total
            self.check_warn(ratio >= 1 / 3,
                            '参考文献', f'外文参考文献应≥1/3 (外文: {foreign}/{total}, {ratio:.0%})')

    # ── 引用标注检查 ──
    def check_citations(self):
        in_refs = False
        for para in self.paragraphs:
            text = get_text(para)
            if not text:
                continue
            if text in ('参考文献', '参 考 文 献'):
                in_refs = True
                continue
            if match_thanks_title(text) or match_appendix(text):
                in_refs = False
                continue
            if in_refs:
                continue
            if match_ref(text):  # 跳过参考文献条目
                continue
            if match_h1(text) or match_h2(text) or match_h3(text):
                if re.search(r'\[\d+\]', text):
                    self.check(False, '引用', f'标题"{text[:20]}"中不应有引用标注')
                continue
            for run in para.runs:
                if re.search(r'\[\d+[\d,\-]*\]', run.text):
                    if not run.font.superscript:
                        if not re.search(r'由文献|见文献', para.text):
                            self.check(False, '引用', f'引用标注"{run.text[:20]}"应为上标格式')

    def _check_ref_entry_mixed_font(self, para, text):
        """检查参考文献条目中英文部分字体"""
        for run in para.runs:
            run_text = run.text.strip()
            if not run_text:
                continue
            has_english = bool(re.search(r'[a-zA-Z]', run_text))
            if has_english:
                rPr = run._element.find(qn('w:rPr'))
                ascii_font = run.font.name
                if rPr is not None:
                    rFonts = rPr.find(qn('w:rFonts'))
                    if rFonts is not None:
                        ascii_font = rFonts.get(qn('w:ascii')) or ascii_font
                if ascii_font and ascii_font != 'Times New Roman':
                    self.check_warn(False, '字体',
                                    f'参考文献中英文部分"{run_text[:20]}"应用Times New Roman (实际: {ascii_font})')

    # ── 表格内文字格式 ──
    def check_table_cell_format(self):
        for ti, table in enumerate(self.doc.tables):
            if ti == 0 and len(table.rows) <= 10:
                first_cell_text = ''
                if table.rows[0].cells:
                    first_cell_text = table.rows[0].cells[0].text.strip()
                if first_cell_text in ('题目', '作者', '学院'):
                    continue
            for ri, row in enumerate(table.rows):
                for ci, cell in enumerate(row.cells):
                    for para in cell.paragraphs:
                        for run in para.runs:
                            if not run.text.strip():
                                continue
                            run_text = run.text.strip()
                            if run.font.size:
                                self.check(abs(run.font.size - SPEC['table_cell']['size']) <= TOLERANCE,
                                           '表格', f'表{ti+1}单元格"{run_text[:15]}"字号应为五号/10.5pt')
                            font_name = run.font.name
                            if font_name:
                                has_chinese = bool(re.search(r'[\u4e00-\u9fff]', run_text))
                                if has_chinese:
                                    self.check(font_name == SPEC['table_cell']['font_cn'],
                                               '表格', f'表{ti+1}中文"{run_text[:15]}"应宋体 (实际: {font_name})')

    # ── 另起页检查 ──
    def check_all_page_breaks(self):
        """检查所有需要另起页的部分"""
        for i, para in enumerate(self.paragraphs):
            text = get_text(para)
            if not text:
                continue
            if self._is_in_toc(i):
                continue
            if self._is_in_cover(i):
                continue

            needs_break = (
                re.match(r'^摘\s*要$', text) or
                text.strip().upper() == 'ABSTRACT' or
                text in ('目 录', '目录') or re.match(r'^目\s+录$', text) or
                match_h1(text) or
                text in ('参考文献', '参 考 文 献') or
                match_thanks_title(text) or
                match_appendix(text)
            )

            if needs_break and i > 0:
                pPr = para._element.find(qn('w:pPr'))
                has_page_break = False
                if pPr is not None:
                    has_page_break = (pPr.find(qn('w:pageBreakBefore')) is not None or
                                      pPr.find(qn('w:sectPr')) is not None)
                if not has_page_break:
                    prev_para = self.paragraphs[i - 1]
                    prev_pPr = prev_para._element.find(qn('w:pPr'))
                    if prev_pPr is not None:
                        has_page_break = prev_pPr.find(qn('w:sectPr')) is not None
                    if not has_page_break:
                        for run_elem in prev_para._element.findall(qn('w:r')):
                            for br in run_elem.findall(qn('w:br')):
                                if br.get(qn('w:type')) == 'page':
                                    has_page_break = True
                                    break
                            if has_page_break:
                                break
                if not has_page_break:
                    section_name = {
                        '摘 要': '中文摘要', '摘要': '中文摘要', 'ABSTRACT': '英文摘要',
                        '目 录': '目录', '目录': '目录',
                        '参考文献': '参考文献', '参 考 文 献': '参考文献',
                        '致 谢': '致谢', '致谢': '致谢',
                    }.get(text, None)
                    if section_name is None:
                        if re.match(r'^摘\s+要$', text):
                            section_name = '中文摘要'
                        elif re.match(r'^目\s+录$', text):
                            section_name = '目录'
                        elif match_h1(text):
                            section_name = '每章'
                        elif match_appendix(text):
                            section_name = '附录'
                    self.check(False, '分页', f'"{section_name or text[:15]}"前应另起页')

    # ── 摘要正文格式 ──
    def check_abstract_body_format(self):
        in_cn = False
        in_en = False
        for para in self.paragraphs:
            text = get_text(para)
            if re.match(r'^摘\s*要$', text):
                in_cn, in_en = True, False
                continue
            if text.strip().upper() == 'ABSTRACT':
                in_cn, in_en = False, True
                continue
            if text.startswith('关键词') or text.startswith('Keywords'):
                in_cn, in_en = False, False
                continue
            if match_h1(text):
                in_cn, in_en = False, False
                continue
            if in_cn and text:
                run = get_run_info(para)
                if run and run.font.size:
                    self.check(abs(run.font.size - SPEC['abstract_body']['size']) <= TOLERANCE,
                               '摘要', f'中文摘要内容"{text[:20]}"字号应为小4号/12pt')
            if in_en and text:
                run = get_run_info(para)
                if run and run.font.size:
                    self.check(abs(run.font.size - SPEC['en_abstract_body']['size']) <= TOLERANCE,
                               '摘要', f'英文摘要内容"{text[:20]}"字号应为小4号/12pt')

    # ── 公式检查 ──
    def check_formulas(self):
        for para in self.paragraphs:
            has_math = False
            for run in para.runs:
                if run._element.findall(qn('w:object')) or run._element.findall(qn('m:oMath')):
                    has_math = True
                    break
            if has_math:
                self.check(para.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER,
                           '公式', '公式应居中')

    # ── 结构完整性 ──
    def check_structure(self):
        all_text = '\n'.join(get_text(p) for p in self.paragraphs)
        self.check('摘要' in all_text or bool(re.search(r'摘\s+要', all_text)),
                   '结构', '应包含中文摘要')
        self.check('ABSTRACT' in all_text.upper(), '结构', '应包含英文摘要')
        self.check('目录' in all_text or bool(re.search(r'目\s+录', all_text)),
                   '结构', '应包含目录')
        self.check('参考文献' in all_text or '参 考 文 献' in all_text,
                   '结构', '应包含参考文献')
        self.check('致谢' in all_text or bool(re.search(r'致\s+谢', all_text)),
                   '结构', '应包含致谢')

    # ── 行距检查（标题单倍行距） ──
    def check_heading_line_spacing(self):
        """检查标题行距：单倍行距"""
        for i, para in enumerate(self.paragraphs):
            text = get_text(para)
            if not text or self._is_in_toc(i) or self._is_in_cover(i):
                continue
            if match_h1(text) or match_h2(text) or match_h3(text) or match_h4(text):
                pf = para.paragraph_format
                if pf.line_spacing is not None:
                    self.check(abs(pf.line_spacing - 1.0) <= 0.05,
                               '行距', f'标题"{text[:20]}"行距应为单倍 (实际: {pf.line_spacing})')

    # ── 汇总 ──
    def run_all(self):
        print(f"\n{'='*60}")
        print(f"  湖南科技大学硕士学位论文格式检查")
        print(f"  文件: {self.filepath}")
        print(f"{'='*60}")

        self.check_page_setup()
        self.check_header()
        self.check_page_number()
        self.check_font_paragraph()
        self.check_heading_line_spacing()
        self.check_citations()
        self.check_formulas()
        self.check_table_cell_format()
        self.check_bilingual_caption()
        self.check_all_page_breaks()
        self.check_structure()
        self.check_keyword_count()
        self.check_paper_title_length()
        self.check_abstract_body_format()
        self.check_abstract_word_count()
        self.check_foreign_ref_ratio()

        # 输出结构化报告
        print(f"\n{'-'*60}")
        print(f"  [OK] 通过: {len(self.passed)} 项")
        if self.errors:
            print(f"  [!!] 错误: {len(self.errors)} 项")
            for e in self.errors:
                print(f"    {e}")
        if self.warnings:
            print(f"  [??] 警告: {len(self.warnings)} 项")
            for w in self.warnings:
                print(f"    {w}")
        print(f"{'-'*60}")

        # 按类别统计
        categories = {}
        for e in self.errors:
            cat = e.split(']')[0].lstrip('[')
            categories.setdefault(cat, {'errors': 0, 'warnings': 0})
            categories[cat]['errors'] += 1
        for w in self.warnings:
            cat = w.split(']')[0].lstrip('[')
            categories.setdefault(cat, {'errors': 0, 'warnings': 0})
            categories[cat]['warnings'] += 1
        if categories:
            print(f"\n  分类汇总:")
            for cat, counts in sorted(categories.items()):
                parts = []
                if counts['errors']:
                    parts.append(f"✗{counts['errors']}")
                if counts['warnings']:
                    parts.append(f"⚠{counts['warnings']}")
                print(f"    {cat}: {' '.join(parts)}")
        print(f"{'─'*60}\n")

        return len(self.errors) == 0


def main():
    if len(sys.argv) < 2:
        print("用法: python check_thesis.py <docx文件路径>")
        print("示例: python check_thesis.py 硕士论文.docx")
        sys.exit(1)

    filepath = sys.argv[1]
    if not os.path.exists(filepath):
        print(f"错误: 文件不存在 - {filepath}")
        sys.exit(1)

    checker = ThesisChecker(filepath)
    ok = checker.run_all()
    sys.exit(0 if ok else 1)


if __name__ == '__main__':
    main()
